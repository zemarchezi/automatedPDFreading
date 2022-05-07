#%%
from PIL import Image
from io import StringIO
import matplotlib.pyplot as plt
from docx import Document
import zipfile
from polyglot.detect import Detector
from generateLatexFile import *
import cv2
import numpy as np
import pandas as pd

# %%
def saveFigs(zipf, zipimgpath, outputpath, figname, crop):
    image1 = zipf.open(zipimgpath).read()
    img = cv2.imdecode(np.frombuffer(image1, np.uint8),3)
    if crop:
        imgCropped = img[crop[0]:crop[1],crop[2]:crop[3]]
    else:
        imgCropped = img
    cv2.imwrite(f'{outputpath}/{figname}.png',imgCropped)
    return f'{outputpath}/{figname}.png'

def get_bold_list(para):
    bold_list = []
    for run in para.runs:
        bold_list.append(run.bold)
    return bold_list


def extractTables(docum):
    tables = list()
    for table in docum.tables:

        data = []

        keys = None
        for i, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)

            if i == 0:
                keys = tuple(text)
                continue
            row_data = dict(zip(keys, text))
            data.append(row_data)
            # print (data)

        df = pd.DataFrame(data)
        tableStr = "\\begin{table}[H]\n\\centering\n \\begin{tabular}{|" + " | ".join(["c"] * len(df.columns)) + "|}\n"
        for i, row in df.iterrows():
            tableStr += "\\hline  \n"
            tableStr += " & ".join([str(x) for x in row.values]) + " \\\\\n"
        tableStr += "\\hline  \n\\end{tabular} \n \\caption{%s} \n \\end{table}"

        tables.append(tableStr)

    return tables

def extractFiguresTextRoti(docPath, filename):

    z = zipfile.ZipFile(docPath+filename)
    all_files = z.namelist()
    image = [x for x in all_files if x.startswith('word/media/')]
    images = []
    for im in image:
        img1 = z.open(im).read()
        img = cv2.imdecode(np.frombuffer(img1, np.uint8),3)
        if img.shape[0]>120:
            images.append(im)

    document = Document(docPath+filename)
    texs = []
    textos = []
    for para in document.paragraphs:
        texs.append(para.text)
        textos.append(para.text)
    removeList = ['Ionos', 'Carol', 'Summa', 'Resum', 'Figu', 'Tab']
    for ii in removeList:
        for nn, jj in enumerate(texs):
            if jj.startswith(ii):
                texs.pop(nn)

    indexRef = [i for i in range(len(texs)) if texs[i].startswith('Refe')]
    # print(indexRef)
    if len(indexRef) > 0:
        texs = texs[0:indexRef[0]]

    tabs = extractTables(document)
    dictsposition = {}
    fig = 0
    tab = 0
    for n, i in enumerate(textos):
        if i.startswith('Fig'):
            fig += 1
            ll = i
            if fig-1 < len(images):
                ims = images[fig-1]
            else:
                ims = images[-1]
            dictsposition[f"Figure_{fig}"] = {'leg': ll,
                                              'fig': ims}
        if i.startswith('Tab'):
            tab+=1
            ll = i
            if tab-1 < len(tabs):
                tbs = tabs[tab-1]
            else:
                tbs = tabs[-1]
            dictsposition[f"Table_{tab}"] = {'leg': i,
                                             'tab': tbs}

    return dictsposition, texs, z


def constructLatexFileRoti(docPath, filename, outputFigure):
    dictsposition, texst, zipf = extractFiguresTextRoti(docPath, filename)

    textpt = '\section{ROTI} \n \subsection{Responsável: Carolina de Sousa do Carmo} \n \n'
    texten = '\section{ROTI} \n \subsection{Responsible: Carolina de Sousa do Carmo} \n \n'

    for i in texst:
        if len(i)>1:
            detector = Detector(i)
            if detector.language.code == 'en':
                texten += i + '\n\n'
            if detector.language.code == 'pt':
                textpt += i + '\n\n'
    
    # if detector.language.code == 'pt':
    #     text = 
    # if detector.language.code == 'en':
    #     text = '\section{ROTI} \n \subsection{Responsible: Carolina de Sousa do Carmo} \n \n'
    keys = list(dictsposition.keys())
    
    includeFigure = """\\begin{figure}[H]\n    
                        \\centering\n   
                             \\includegraphics[width=14cm]{./%s}\n
                             \\caption{%s}
                        \\end{figure}\n
                     """

    for i in range(len(keys)):
            
        if keys[i].startswith("Tab"):
            tbs = dictsposition[keys[i]]['tab'] % (dictsposition[keys[i]]['leg'].split('–')[-1])
            if dictsposition[keys[i]]['leg'].startswith('Table'):
                texten += tbs
            if dictsposition[keys[i]]['leg'].startswith('Tabela'):
                textpt += tbs

        if keys[i].startswith("Fig"):

            outfigpath = saveFigs(zipf, dictsposition[keys[i]]['fig'], f"{outputFigure}", f'figureROTI_{i}', crop=False)
            
            if dictsposition[keys[i]]['leg'].startswith('Figure'):
                texten += includeFigure % ('/'.join(outfigpath.split('/')[2:]), dictsposition[keys[i]]['leg'].split('–')[-1])
            if dictsposition[keys[i]]['leg'].startswith('Figura'):
                textpt += includeFigure % ('/'.join(outfigpath.split('/')[2:]), dictsposition[keys[i]]['leg'].split('–')[-1])

    # textSum = '\n\n'.join(texst[dictsposition['locSummary']+1:])
    # text += textSum


    return textpt, texten



PATH = './data/'

filename = 'Resumo_semana2207_Carolina.docx'

#%%
ssa = extractFiguresTextRoti(PATH, filename)
textpt, texten = constructLatexFileRoti(PATH, filename, outputFigure='./latexText/figures')

generateLaTexFile(textpt, EnPt=False, outputPath='./latexText', date='25/04/2022')
